# -*- coding:utf-8 -*-
  
# 此程序可扫描Log中的docx文件并返回基本信息
  
import docx
from docx import Document
  
test_d = '../log/sublime搭建python的集成开发环境.docx'
  
def docxInfo(addr):
 document = Document(addr)
  
 info = {'title':[],
 'keywords':[],
 'author':[],
 'date':[],
 'question':[]}
  
 lines = [0 for i in range(len(document.paragraphs))]
 k = 0
 for paragraph in document.paragraphs:
     lines[k] = paragraph.text
     k = k+1
  
 index = [0 for i in range(5)]
 k = 0
 for line in lines:
     if line.startswith('标题'):
        index[0] = k
     if line.startswith('关键词'):
        index[1] = k
     if line.startswith('作者'):
        index[2] = k
     if line.startswith('日期'):
        index[3] = k
     if line.startswith('问题描述'):
        index[4] = k
     k = k+1
  
 info['title'] = lines[index[0]+1]
  
 keywords = []
 for line in lines[index[1]+1:index[2]]:
     keywords.append(line)
     info['keywords'] = keywords
      
     info['author'] = lines[index[2]+1]
      
     info['date'] = lines[index[3]+1]
      
     info['question'] = lines[index[4]+1]
  
 return info
  
if __name__ == '__main__':
 print(docxInfo(test_d))